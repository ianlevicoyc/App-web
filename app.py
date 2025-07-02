from flask import Flask, request, jsonify, redirect, url_for, render_template, send_from_directory, flash, session, send_file, Response, current_app
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from flask_bcrypt import Bcrypt
from flask_cors import CORS
from models import db, User, Image, ExcelFile, Graph
import os
import cohere
from datetime import datetime
from dotenv import load_dotenv
import re
from itertools import cycle
from sqlalchemy.orm import joinedload
from flask import g
from flask import send_file
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
import io
from flask_migrate import Migrate
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import Boolean, Column
from werkzeug.utils import secure_filename
import pandas as pd
import matplotlib.pyplot as plt
import tkinter
import math
from werkzeug.utils import secure_filename
import numpy as np
from matplotlib import dates as mdates
from urllib.parse import unquote
from models import db, SupportMessage
from functools import wraps
import hashlib
import joblib
from pytz import timezone
from datetime import datetime
from sklearn.ensemble import RandomForestClassifier
from sklearn.model_selection import train_test_split
from sklearn.metrics import accuracy_score
import pandas as pd
from sklearn.linear_model import LinearRegression
from flask_mail import Mail, Message
import pytz
import joblib
# Cargar variables de entorno

#.\venv\Scripts\Activate


load_dotenv()

app = Flask(__name__)
CORS(app)


tkinter.Tk().withdraw()  # Esto asegura que Tk no se muestre ni interfiera.
os.environ['MPLBACKEND'] = 'Agg'

# Configuración de clave secreta para sesiones y base de datos
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///proyecto.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db.init_app(app)
migrate = Migrate(app, db)

app.secret_key = 'Qwerty123456'



# Configurar Flask-Mail
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USE_SSL'] = False
app.config['MAIL_USERNAME'] = 'supervisorilevicoy@gmail.com'  # Dirección de Gmail
app.config['MAIL_PASSWORD'] = 'ykncqbdvfywmkobj' 
app.config['MAIL_DEFAULT_SENDER'] = 'supervisorilevicoy@gmail.com'


mail = Mail(app)

# Extensiones permitidas para Excel
ALLOWED_EXTENSIONS = {'xls', 'xlsx'}
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS
# Configuración para subir archivos Excel
EXCEL_UPLOAD_FOLDER = './uploads/excel'
app.config['EXCEL_UPLOAD_FOLDER'] = EXCEL_UPLOAD_FOLDER

app.config['IMAGE_UPLOAD_FOLDER'] = './uploads/images'
app.config['EXCEL_UPLOAD_FOLDER'] = './uploads/excel'


# Crear carpeta si no existe
if not os.path.exists(EXCEL_UPLOAD_FOLDER):
    os.makedirs(EXCEL_UPLOAD_FOLDER)

UPLOAD_FOLDER = os.path.abspath('./uploads/images')
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
    print(f"Carpeta creada: {UPLOAD_FOLDER}")



def enviar_notificacion_email(destinatario, asunto, cuerpo):
    try:
        mensaje = Message(asunto, recipients=[destinatario])
        mensaje.body = cuerpo
        mail.send(mensaje)
        print(f"Correo enviado a {destinatario}")
    except Exception as e:
        print(f"Error al enviar correo: {e}")




def generar_grafico_profesional(data, filename, units_filter, title, category):
    """
    Genera un gráfico profesional filtrando los datos cada hora y seleccionando el valor más alto sin aproximaciones.
    Añade una línea de tendencia desde el primer hasta el último punto.
    Guarda la información de categoría para identificar el tipo de gráfico.
    """
    try:
        # Definir colores para las unidades
        colores = {
            "Aceleración": "#E1B800",  # Amarillo oscuro
            "Velocidad": "blue",       # Azul
            "Envolvente de Aceleración": "red"  # Rojo
        }

        # Mapear título a color
        color = colores.get(title, "gray")

        # Filtrar por unidad específica
        filtered_data = data[data['units'] == units_filter]
        print(f"[DEBUG] Datos filtrados (units == {units_filter}): {len(filtered_data)} filas")

        if filtered_data.empty:
            print(f"[INFO] No hay datos para Units = {units_filter}")
            return None, None  # Retorna None si no hay datos

        # Convertir 'data_time' a datetime
        filtered_data['data_time'] = pd.to_datetime(filtered_data['data_time'], errors='coerce')
        filtered_data = filtered_data.dropna(subset=['data_time', 'realvalue'])

        if filtered_data.empty:
            print(f"[INFO] Datos filtrados están vacíos después de convertir 'data_time' y eliminar NaN.")
            return None, None

        # Agrupar por hora y conservar el valor real más alto sin alterarlo
        grouped_data = filtered_data.groupby(filtered_data['data_time'].dt.floor('1H'))['realvalue'].max()

        if grouped_data.empty:
            print(f"[INFO] Datos agrupados están vacíos después de agrupar por hora.")
            return None, None

        # Crear el gráfico
        plt.style.use('ggplot')
        plt.figure(figsize=(18, 8))  # Gráfico más grande
        plt.plot(grouped_data.index, grouped_data, label=title, color=color, linewidth=2, marker='o')

        # Ajustar los límites del eje X
        plt.xlim(grouped_data.index[0], grouped_data.index[-1])

        # Línea de tendencia
        x = np.array([0, len(grouped_data) - 1])  # Índices de los puntos inicial y final
        y = np.array([grouped_data.iloc[0], grouped_data.iloc[-1]])  # Valores inicial y final

        # Calcular la ecuación de la línea (pendiente e intersección)
        pendiente = (y[1] - y[0]) / (x[1] - x[0])
        interseccion = y[0] - pendiente * x[0]
        tendencia_x = np.arange(len(grouped_data))  # Todos los índices del rango
        tendencia_y = pendiente * tendencia_x + interseccion  # Valores de la línea de tendencia

        # Graficar la línea de tendencia
        plt.plot(grouped_data.index, tendencia_y, label="Tendencia", color="green", linestyle="--", linewidth=2)

        # Configuración de título y etiquetas
        plt.title(f"{title} - Valor Máximo por Hora (Con Línea de Tendencia)")
        plt.xlabel("Tiempo (1 Hora)")
        plt.ylabel("Valor Máximo")
        plt.legend()
        plt.grid(True)

        # Guardar el gráfico con un nombre único
        graph_folder = os.path.join('static', 'graphs')
        os.makedirs(graph_folder, exist_ok=True)  # Crear la carpeta si no existe

        # Generar un nombre único para el archivo usando secure_filename y timestamp
        timestamp = datetime.utcnow().strftime("%Y%m%d%H%M%S")
        graph_filename = secure_filename(f"{filename.rsplit('.', 1)[0]}_{category}_Units_{units_filter}_MaxHour_{timestamp}.png")
        graph_path = os.path.join(graph_folder, graph_filename)

        plt.savefig(graph_path, dpi=300)
        plt.close()

        print(f"[INFO] Gráfico generado y guardado en: {graph_path}")

        # VALIDAR SI YA EXISTE EN LA BASE DE DATOS
        existing_graph = Graph.query.filter_by(graph_path=f"graphs/{graph_filename}").first()
        if existing_graph:
            print(f"[DEBUG] Gráfico ya existe en la base de datos: {graph_path}")
            return f"graphs/{graph_filename}", category

        print(f"[DEBUG] Nuevo gráfico agregado: {graph_path}")
        return f"graphs/{graph_filename}", category  # Devuelve la ruta relativa y la categoría

    except Exception as e:
        print(f"[ERROR] Error al generar el gráfico: {e}")
        return None, None






def normalize_columns(dataframe):
    """
    Normaliza las columnas de un DataFrame:
    - Convierte a minúsculas.
    - Elimina espacios en blanco antes y después del nombre.
    - Reemplaza espacios internos con '_'.
    """
    dataframe.columns = dataframe.columns.str.strip().str.lower().str.replace(' ', '_')
    return dataframe


def clasificar_ml(row):
    """
    Clasifica utilizando el modelo preentrenado.
    """
    try:
        if modelo_ml:
            # Crear un DataFrame con los nombres de columnas esperados
            feature_names = ['realvalue', 'units', 'pointindex', 'axis']
            features = pd.DataFrame([[row['realvalue'], row['units'], row['pointindex'], row['axis']]],
                                    columns=feature_names)
            prediction = modelo_ml.predict(features)
            return prediction[0]
        else:
            return "Error en ML"
    except Exception as e:
        print(f"[ERROR] Error en clasificación ML: {e}")
        return "Error en ML"


def cargar_modelo_ml():
    global modelo_ml
    if modelo_ml is None:
        try:
            modelo_ml = joblib.load("model_anomalias.joblib")
            (
                r"C:\Users\ianle\Desktop\Trabajos inacap\Octavo SEMESTRE\Prototipo proyecto de titulo RESPALDO DE SEGURIDAD 1\Backend\model.joblib"
            )
            print("[INFO] Modelo de Machine Learning cargado exitosamente.")
        except Exception as e:
            print(f"[ERROR] No se pudo cargar el modelo de ML: {e}")
            modelo_ml = None

modelo_ml = None
cargar_modelo_ml()


@app.route('/upload_excel', methods=['GET', 'POST'])
@login_required
def upload_excel():
    if request.method == 'POST':
        try:
            cargar_modelo_ml()  # Cargar el modelo si aún no está disponible
            modelo_anomalias = joblib.load("model_anomalias.joblib")  # Cargar el nuevo modelo de anomalías

            # Verificar archivo
            if 'file' not in request.files or request.files['file'].filename == '':
                flash("No se seleccionó ningún archivo válido.", "danger")
                return redirect(url_for('upload_excel'))

            file = request.files['file']
            if not allowed_file(file.filename):
                flash("Formato de archivo no permitido.", "danger")
                return redirect(url_for('upload_excel'))

            # Guardar archivo
            filepath = os.path.join(app.config['EXCEL_UPLOAD_FOLDER'], secure_filename(file.filename))
            file.save(filepath)
            print(f"[DEBUG] Archivo guardado en: {filepath}")

            # Leer y limpiar datos
            data = pd.read_excel(filepath)
            data = normalize_columns(data)

            # Validar columnas
            required_columns = {'data_time', 'realvalue', 'units', 'pointindex', 'axis'}
            if not required_columns.issubset(data.columns):
                flash("El archivo debe contener las columnas requeridas.", "danger")
                return redirect(url_for('upload_excel'))

            # Limpieza de datos
            data['realvalue'] = pd.to_numeric(data['realvalue'], errors='coerce')
            data['data_time'] = pd.to_datetime(data['data_time'], errors='coerce')
            data = data.dropna(subset=['realvalue', 'data_time'])

            # Filtrar datos
            pointindex_filter = int(request.form.get('point_index'))
            axis_filter = int(request.form.get('axis'))
            filtered_data = data[
                (data['pointindex'] == pointindex_filter) & 
                (data['axis'] == axis_filter)
            ].copy()

            # Mapear los valores de punto y eje
            point_mapping = {
                "1": "Motor Lado Libre",
                "2": "Motor Lado Reductor",
                "3": "Reductor Entrada De Alta",
                "4": "Reductor Salida",
                "5": "Contraeje Lado Motor",
                "6": "Contraeje Lado Molino",
            }
            axis_mapping = {
                "1": "Horizontal",
                "2": "Vertical",
                "3": "Axial",
            }

            selected_point = point_mapping.get(str(pointindex_filter), "No especificado")
            selected_axis = axis_mapping.get(str(axis_filter), "No especificado")

            # Determinar las fechas de inicio y fin del período de los datos
            start_date = filtered_data['data_time'].min().date()
            end_date = filtered_data['data_time'].max().date()

            # Crear registro en ExcelFile
            new_file = ExcelFile(
                filename=file.filename,
                classification="Análisis Generado",
                uploaded_by=current_user.id
            )
            db.session.add(new_file)
            db.session.commit()  # Commit inicial para obtener el ID
            print(f"[DEBUG] Registro creado en ExcelFile con ID: {new_file.id}")

            # Generar gráficos y guardarlos en la base de datos
            graph_acc, category_acc = generar_grafico_profesional(filtered_data, file.filename, 0, "Aceleración", "Aceleracion")
            graph_vel, category_vel = generar_grafico_profesional(filtered_data, file.filename, 2, "Velocidad", "Velocidad")
            graph_env, category_env = generar_grafico_profesional(filtered_data, file.filename, 6, "Envolvente de Aceleración", "Envolvente")

            # Verificar y guardar gráficos
            for graph_path, category in [(graph_acc, category_acc), (graph_vel, category_vel), (graph_env, category_env)]:
                if graph_path:  # Verificar si el gráfico fue generado correctamente
                    existing_graph = Graph.query.filter_by(graph_path=graph_path, excel_file_id=new_file.id).first()
                    if not existing_graph:
                        new_graph = Graph(
                            graph_path=graph_path,
                            category=category,
                            excel_file_id=new_file.id,
                            uploaded_at=datetime.utcnow(),
                            start_date=start_date,  # Guardar fecha de inicio
                            end_date=end_date      # Guardar fecha de fin
                        )
                        db.session.add(new_graph)

            db.session.commit()  # Confirmar todos los cambios en la base de datos
            print(f"[DEBUG] Todos los gráficos se guardaron en la base de datos.")

            # Clasificación ML
            filtered_data['ml_classification'] = filtered_data.apply(clasificar_ml, axis=1)

            # Clasificación de anomalías (separadas por units)
            filtered_data['anomalías'] = filtered_data.apply(clasificar_anomalias, axis=1)
            anomaly_summary = {
                "Aceleración": filtered_data[filtered_data['units'] == 0]['anomalías'].value_counts().to_dict(),
                "Velocidad": filtered_data[filtered_data['units'] == 2]['anomalías'].value_counts().to_dict(),
                "Envolvente de Aceleración": filtered_data[filtered_data['units'] == 6]['anomalías'].value_counts().to_dict(),
            }

            # Análisis ML adicional usando el modelo de anomalías
            def predecir_anomalias(row):
                try:
                    features = pd.DataFrame([[row['realvalue'], row['units'], row['pointindex'], row['axis']]], 
                                            columns=['realvalue', 'units', 'pointindex', 'axis'])
                    prediction = modelo_anomalias.predict(features)
                    return "Anomalía Detectada" if prediction[0] == 1 else "Normal"
                except Exception as e:
                    print(f"[ERROR] Predicción de anomalías: {e}")
                    return "Error"

            # Generar predicciones y corregir copia
            filtered_data['prediccion_anomalias'] = filtered_data.apply(predecir_anomalias, axis=1)

            # Enviar notificación al supervisor (correo fijo)
            # Enviar notificación al supervisor (correo fijo)
            supervisor_email = 'supervisorilevicoy@gmail.com'
            asunto = "📄 [Nuevo Informe Generado] Análisis Técnico Disponibles"

            cuerpo = f"""
            Estimado/a Supervisor/a,

            Nos complace informarle que un nuevo informe técnico ha sido generado y está disponible en el sistema para su revisión.

                Detalles del Informe:
            - Nombre del Archivo: {file.filename}
            - Generado por: {current_user.role.capitalize()} {current_user.username}
            - Fecha de Generación: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
            - Punto de Medición: {selected_point}
            - Eje Medido: {selected_axis}

            Por favor, acceda al sistema para revisar el análisis técnico y los gráficos generados.

            Atentamente,  
            **Equipo de Inspecciones Chile SpA**

            ---
            Este es un mensaje generado automáticamente. Por favor, no responda a este correo.
            """

            # Enviar correo
            enviar_notificacion_email(supervisor_email, asunto, cuerpo)


            # Convertir a lista de diccionarios para evitar problemas con Jinja2
            predicciones_finales = filtered_data[['data_time', 'realvalue', 'prediccion_anomalias']].to_dict(orient='records')

            # Preparar información para generar el informe
            datos_generales = {
                "nombre_archivo": file.filename,
                "usuario": f"{current_user.role} {current_user.username}",
                "fecha": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                "punto_medicion": selected_point,
                "eje_medido": selected_axis
            }

            clasificacion_ml = filtered_data['ml_classification'].value_counts().to_dict()

            graficos = [
                {"label": "Aceleración", "path": graph_acc},
                {"label": "Velocidad", "path": graph_vel},
                {"label": "Envolvente de Aceleración", "path": graph_env},
            ]

            # Generar el informe con datos más completos
            informe_texto = generar_texto_informe(
                datos_generales,
                clasificacion_ml,
                "Estable",  # Ajustar según la tendencia observada
                anomaly_summary
            )

            # Almacenar en sesión
            session['ml_summary'] = clasificacion_ml
            session['anomaly_summary'] = anomaly_summary
            session['graphs'] = graficos
            session['informe_texto'] = informe_texto
            session['punto_medicion'] = selected_point
            session['eje_medido'] = selected_axis

            print("[DEBUG] Datos almacenados en la sesión correctamente.")

            flash("Archivo procesado y notificación enviada al supervisor.", "success")

            return render_template(
                "informe_excel_template.html",
                graphs=session['graphs'],
                filename=file.filename,
                usuario=f"{current_user.role} {current_user.username}",
                fecha=datetime.now(),
                ml_summary=session['ml_summary'],
                anomaly_summary=session['anomaly_summary'],
                predicciones=predicciones_finales,
                informe_texto=session['informe_texto'],
                punto_medicion=selected_point,
                eje_medido=selected_axis
            )

        except Exception as e:
            print(f"[ERROR] Error al procesar el archivo: {e}")
            flash(f"Error al procesar el archivo: {e}", "danger")
            return redirect(url_for('upload_excel'))

    return render_template('upload_excel.html')






def clasificar_anomalias(row):
    """
    Clasifica las anomalías según units:
    - Units == 0: Aceleración
    - Units == 2: Velocidad
    - Units == 6: Envolvente de Aceleración
    """
    gE = row['realvalue']
    units = row['units']

    if units == 0:  # Aceleración
        if gE <= 0.1:
            return "Buena"
        elif gE <= 0.5:
            return "Satisfactoria"
        elif gE <= 1:
            return "No Satisfactoria"
        else:
            return "Inaceptable"

    elif units == 2:  # Velocidad
        if gE <= 0.71:
            return "Buena"
        elif gE <= 1.8:
            return "Satisfactoria"
        elif gE <= 4.5:
            return "No Satisfactoria"
        else:
            return "Inaceptable"

    elif units == 6:  # Envolvente de Aceleración
        if gE <= 0.1:
            return "Buena"
        elif gE <= 0.5:
            return "Satisfactoria"
        elif gE <= 1:
            return "No Satisfactoria"
        else:
            return "Inaceptable"

    return "Desconocido"



@app.route('/enable_user/<int:user_id>', methods=['POST'])
@login_required
def enable_user(user_id):
    if g.user_role == 'admin':
        user = User.query.get_or_404(user_id)
        user.is_active = True  # Activar al usuario
        db.session.commit()
        flash('Usuario habilitado correctamente.', 'success')
    else:
        flash('No tienes permisos para realizar esta acción.', 'danger')
    return redirect(url_for('admin_area'))


@app.route('/disable_user/<int:user_id>', methods=['POST'])
@login_required
def disable_user(user_id):
    user = User.query.get_or_404(user_id)
    user.is_active = False  # Desactivar al usuario
    db.session.commit()
    flash('Usuario deshabilitado con éxito.', 'success')
    return redirect(url_for('admin_area'))


def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if session.get('user_role') != 'admin':
            return redirect(url_for('index'))  # Redirige si no es admin
        return f(*args, **kwargs)
    return decorated_function

@app.route('/support', methods=['POST'])
def support():
    data = request.get_json()
    email = data.get('email')
    message = data.get('message')

    if not email or not message:
        return jsonify({'success': False, 'error': 'Email y mensaje son obligatorios'}), 400

    # Guardar el mensaje en la base de datos
    new_message = SupportMessage(email=email, message=message)
    db.session.add(new_message)
    db.session.commit()

    return jsonify({'success': True, 'message': 'Mensaje enviado correctamente'}), 200


@app.route('/admin/support-messages', methods=['GET'])
def view_support_messages():
    # Obtén los mensajes de soporte desde la base de datos
    messages = SupportMessage.query.order_by(SupportMessage.timestamp.desc()).all()
    return render_template('admin_support_messages.html', messages=messages)


@app.route('/supervisor_dashboard', methods=['GET', 'POST'])
@login_required
def supervisor_dashboard():
    """
    Dashboard del supervisor para mostrar gráficos de Aceleración, Velocidad y Envolvente,
    con filtrado por rango de fechas de los gráficos trabajados.
    """
    start_date = request.form.get('start_date')  # Fecha de inicio desde el formulario
    end_date = request.form.get('end_date')      # Fecha de fin desde el formulario

    try:
        # Configurar la zona horaria de Chile
        chile_tz = timezone('America/Santiago')

        # Consulta base: Join entre Graph y ExcelFile
        query = db.session.query(Graph, ExcelFile).join(ExcelFile, Graph.excel_file_id == ExcelFile.id)

        # Filtrar por rango de fechas si se especifica
        if start_date and end_date:
            start_date_obj = datetime.strptime(start_date, '%Y-%m-%d').date()
            end_date_obj = datetime.strptime(end_date, '%Y-%m-%d').date()
            query = query.filter(Graph.start_date >= start_date_obj, Graph.end_date <= end_date_obj)

        # Obtener resultados de la consulta
        graph_results = query.all()

        # Clasificar gráficos en categorías
        graphs_by_type = {'Aceleración': [], 'Velocidad': [], 'Envolvente de Aceleración': []}

        # Procesar los resultados
        for graph, excel_file in graph_results:
            graph_full_path = os.path.join('static', graph.graph_path)
            if not os.path.exists(graph_full_path):
                print(f"[WARNING] Archivo no encontrado: {graph.graph_path}")
                continue  # Omitir el registro si el archivo no existe

            # Convertir la fecha de subida a la zona horaria de Chile
            uploaded_at = excel_file.upload_date.astimezone(chile_tz)
            print(f"[DEBUG] Fecha original en UTC: {excel_file.upload_date}, convertida a Chile: {uploaded_at}")

            # Lógica para asignar el punto de medición y eje medido (dinámicos o estáticos)
            measurement_point = "Reductor Salida"  # Valor fijo o calculado
            measured_axis = "Horizontal"          # Valor fijo o calculado

            # Información del gráfico
            graph_info = {
                'filename': excel_file.filename,
                'uploaded_by': User.query.get(excel_file.uploaded_by).username,
                'uploaded_at': uploaded_at.strftime('%Y-%m-%d %H:%M:%S'),  # Hora convertida
                'start_date': graph.start_date.strftime('%Y-%m-%d'),
                'end_date': graph.end_date.strftime('%Y-%m-%d'),
                'graph_path': graph.graph_path,
                'measurement_point': measurement_point,  # Campo dinámico
                'measured_axis': measured_axis           # Campo dinámico
            }

            # Clasificar según la categoría
            if graph.category.lower() == "aceleracion":
                graphs_by_type['Aceleración'].append(graph_info)
            elif graph.category.lower() == "velocidad":
                graphs_by_type['Velocidad'].append(graph_info)
            elif graph.category.lower() == "envolvente":
                graphs_by_type['Envolvente de Aceleración'].append(graph_info)

        print("[DEBUG] Gráficos clasificados por tipo:", graphs_by_type)

        return render_template(
            'supervisor_dashboard.html',
            graphs=graphs_by_type,
            start_date=start_date,
            end_date=end_date
        )

    except Exception as e:
        print(f"[ERROR] Error al cargar el dashboard del supervisor: {e}")
        flash(f"Error al cargar el dashboard: {e}", "danger")
        return redirect(url_for('index'))







# Inicializar Bcrypt para contraseñas y Flask-Login para autenticación
bcrypt = Bcrypt(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'

# Cargar el usuario actual
@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# Crear las tablas cuando arranque la aplicación
with app.app_context():
    db.create_all()

# Configuración de la carpeta de subida
UPLOAD_FOLDER = './uploads/images'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Crear la carpeta de imágenes si no existe
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# Ruta para servir imágenes desde la carpeta de uploads
@app.route('/uploads/<filename>')
@login_required
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

# Ruta de subida de archivos (solo carga)
@app.route('/upload', methods=['GET', 'POST'])
@login_required
def upload_file():
    if request.method == 'POST':
        if 'image-file' not in request.files:
            flash("No se seleccionó ningún archivo.", "danger")
            return redirect(url_for('upload_file'))
        
        file = request.files['image-file']
        if file.filename == '':
            flash("Nombre de archivo vacío.", "danger")
            return redirect(url_for('upload_file'))
        
        if file and file.filename.endswith(('.png', '.jpg', '.jpeg')):
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
            file.save(filepath)
            
            # Guardar la imagen en la base de datos con referencia al usuario actual
            new_image = Image(filename=file.filename, uploaded_by=current_user.id)
            db.session.add(new_image)
            db.session.commit()
            
            flash("Imagen subida exitosamente.", "success")
            return redirect(url_for('inspector_dashboard'))
        else:
            flash("Formato de archivo no permitido.", "danger")
            return redirect(url_for('upload_file'))
    
    return render_template('upload_file.html')


@app.before_request
def load_user_role():
    if current_user.is_authenticated:
        g.user_role = current_user.role
        print(f"Rol actual del usuario: {g.user_role}")  # Depuración
    else:
        g.user_role = None



# Configuración de la API de Cohere
cohere_api_key = os.getenv("COHERE_API_KEY")  # Asegúrate de tener esta clave en tu archivo .env
co = cohere.Client(cohere_api_key)



# Función para generar el texto del informe usando Cohere
def generar_texto_informe(nombre_grafico, clasificacion, tendencia, anomaly_summary):
    """
    Genera el texto del informe técnico basado en los datos proporcionados.
    """
    prompt = f"""
    CONTEXTO: Genera un informe técnico profesional y coherente en español basado en los datos proporcionados. Estructura el texto en párrafos claros y concisos, y evita el uso de listas o puntos. Cada sección debe ser un párrafo que explique de manera fluida los resultados y observaciones.

    DATOS:
    - Nombre del gráfico: {nombre_grafico}
    - Clasificación del análisis: {clasificacion}
    - Tendencia observada: {tendencia}
    - Resumen de anomalías: {anomaly_summary}

    ESTRUCTURA DEL INFORME:
    1. Introducción: Presenta el propósito del análisis y el contexto general.
    2. Análisis de resultados: Explica las clasificaciones obtenidas y los hallazgos clave.
    3. Observaciones sobre las tendencias: Describe las tendencias vistas en los gráficos y su impacto.
    4. Recomendaciones: Proporciona acciones concretas basadas en los hallazgos.
    5. Conclusión: Resume los puntos clave y sugiere pasos finales.

    FORMATO:
    - El informe debe estar completamente en español.
    - Usa un lenguaje técnico y directo.
    - Estructura la información en párrafos completos.
    - Evita especulaciones vagas o subjetivas.
    - Evita cualquier tipo de saludo y evita palabras en ingles, solo español.
    - Evita cualquier tipo de agradecimientos.
    - Estructura la información en párrafos completos y fluidos.
    - Usa un lenguaje técnico y directo, evitando términos ambiguos o informales.
    - No inventes datos ni especules sobre información no proporcionada.
    """
    try:
        response = co.generate(
            model='command-xlarge',
            prompt=prompt,
            max_tokens=1700,
            temperature=0.4,
            stop_sequences=["--"],
        )
        # Validar y limpiar la respuesta generada
        texto_generado = response.generations[0].text.strip()
        if not texto_generado:
            raise ValueError("El modelo no generó una respuesta válida.")
        return texto_generado
    except Exception as e:
        print(f"Error generando texto con Cohere: {e}")
        return "Error al generar el informe."






def calcular_dv(rut):
    """
    Calcula el dígito verificador (DV) de un RUT chileno según el algoritmo módulo 11.
    """
    try:
        suma = 0
        multiplicador = 2  # Iniciar con el multiplicador en 2

        print(f"Calculando DV para RUT: {rut}")

        # Procesar cada dígito desde el último al primero
        for digito in reversed(str(rut)):
            producto = int(digito) * multiplicador
            print(f"Dígito: {digito}, Multiplicador: {multiplicador}, Producto: {producto}")
            suma += producto
            multiplicador += 1  # Incrementar el multiplicador
            if multiplicador > 7:  # Reiniciar a 2 después de 7
                multiplicador = 2

        print(f"Suma total: {suma}")

        resto = suma % 11
        dv = 11 - resto

        if dv == 11:
            return '0'
        elif dv == 10:
            return 'K'
        return str(dv)
    except ValueError as e:
        print(f"Error en calcular_dv: {e}")
        return None



def validar_rut(rut, dv):
    """
    Valida un RUT comparando su dígito verificador (DV) calculado con el ingresado.
    """
    try:
        dv_calculado = calcular_dv(rut)
        print(f"Validando RUT: {rut}, DV calculado: {dv_calculado}, DV ingresado: {dv.upper()}")
        return dv_calculado == dv.upper()
    except Exception as e:
        print(f"Error en validar_rut: {e}")
        return False


@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        try:
            # Recibir los datos del formulario
            username = request.form['username']
            email = request.form['email']
            rut = request.form['rut']
            dv = request.form['dv']
            password = request.form['password']
            role = request.form['role']

            # Normalizar el RUT
            rut = rut.replace('-', '').replace('.', '').strip()
            dv = dv.upper().strip()
            rut_completo = f"{rut}-{dv}"

            # Validar el RUT
            if not validar_rut(rut, dv):
                flash("RUT inválido. Por favor, ingresa un RUT correcto.", "danger")
                return render_template('register.html', 
                                    username=username, 
                                    email=email, 
                                    role=role)

            # Validar contraseña
            if len(password) < 12:
                flash("La contraseña debe tener al menos 12 caracteres.", "danger")
                return render_template('register.html', 
                                    username=username, 
                                    email=email, 
                                    rut=rut_completo, 
                                    role=role)

            # Hash de la contraseña
            hashed_password = bcrypt.generate_password_hash(password).decode('utf-8')

            # Crear el usuario
            new_user = User(
                username=username,
                email=email,
                rut=rut_completo,
                password=hashed_password,
                role=role
            )
            db.session.add(new_user)
            db.session.commit()

            flash("Usuario registrado exitosamente.", "success")
            return redirect(url_for('login'))
        
        except Exception as e:
            # Manejar excepciones de la base de datos u otros errores
            print(f"Error al registrar usuario: {e}")
            flash("Hubo un error al registrar el usuario. Por favor, intenta nuevamente.", "danger")
            return render_template('register.html')

    return render_template('register.html')


# Ruta de login
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        # Buscar al usuario en la base de datos usando el nombre de usuario o correo electrónico
        user = User.query.filter_by(username=request.form['username']).first()
        
        # Verificar si el usuario existe y la contraseña es correcta
        if user and bcrypt.check_password_hash(user.password, request.form['password']):
            
            # Verificar si la cuenta está activa
            if not user.is_active:
                flash("Tu cuenta está deshabilitada. Contacta al administrador.", "danger")
                return render_template('login.html', username=request.form['username'])
            
            # Si pasa todas las validaciones, loguear al usuario
            login_user(user)
            flash(f"Bienvenido, {user.username}!", "success")
            
            # Redirigir basado en el rol del usuario
            if user.role == 'admin':
                return redirect(url_for('admin_area'))
            elif user.role == 'supervisor':
                return redirect(url_for('supervisor_dashboard'))
            elif user.role == 'inspector':
                return redirect(url_for('inspector_dashboard'))
            else:
                flash("Rol desconocido. Contacta al administrador.", "danger")
                return redirect(url_for('login'))
        else:
            flash("Nombre de usuario o contraseña incorrectos.", "danger")
    
    return render_template('login.html')




@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

@login_manager.unauthorized_handler
def unauthorized_callback():
    flash("Por favor, inicia sesión para acceder a esta página.", "danger")
    return redirect(url_for('login'))


# Ruta protegida para el área de administradores
@app.route('/admin_area')
@login_required
def admin_area():
    if g.user_role != 'admin':
        # Evitar mostrar mensajes innecesarios para supervisores u otros roles
        return redirect(url_for('index'))
    
    users = User.query.all()
    return render_template('admin_area.html', users=users)


@app.route('/delete_user/<int:user_id>', methods=['POST'])
@login_required
def delete_user(user_id):
    if g.user_role == 'admin':
        user = User.query.get_or_404(user_id)
        if user:
            # Eliminar todas las imágenes asociadas al usuario
            Image.query.filter_by(uploaded_by=user.id).delete()
            db.session.commit()
            db.session.delete(user)
            db.session.commit()
            flash('Usuario eliminado correctamente.', 'success')
        else:
            flash('Usuario no encontrado.', 'danger')
    return redirect(url_for('admin_area'))



@app.route('/edit_user/<int:user_id>', methods=['GET', 'POST'])
@login_required
def edit_user(user_id):
    if current_user.role != 'admin':
        flash("Acceso denegado: Solo los Administradores pueden editar usuarios", "danger")
        return redirect(url_for('admin_area'))
    
    user = User.query.get(user_id)
    if not user:
        flash("Usuario no encontrado.", "danger")
        return redirect(url_for('admin_area'))

    if request.method == 'POST':
        new_username = request.form.get('username')
        new_email = request.form.get('email')
        new_role = request.form.get('role')
        
        user.username = new_username
        user.email = new_email
        user.role = new_role
        db.session.commit()
        
        flash("Usuario actualizado con éxito.", "success")
        return redirect(url_for('admin_area'))
    
    return render_template('edit_user.html', user=user)

# Dashboard de Inspector con paginación
IMAGES_PER_PAGE = 12

@app.route('/inspector_dashboard', methods=['GET'])
@login_required
def inspector_dashboard():
    if current_user.role not in ['inspector', 'supervisor']:
        return "Acceso denegado: Solo los inspectores y supervisores pueden acceder a esta área", 403

    view = request.args.get('view', 'own')  # Por defecto, muestra imágenes propias
    username = request.args.get('username')  # Para filtrar por usuario en la vista general

    IMAGES_PER_PAGE = 12  # Configuración para paginación
    page = int(request.args.get('page', 1))

    # Filtrado según la vista seleccionada
    if view == 'own':  # Mostrar imágenes propias del usuario actual
        images_query = Image.query.filter_by(uploaded_by=current_user.id)
    elif view == 'all':  # Mostrar todas las imágenes
        images_query = Image.query
        if username:  # Si se proporciona un filtro por nombre de usuario
            user = User.query.filter_by(username=username).first()
            if not user:
                flash("Usuario no encontrado.", "danger")
                return redirect(url_for('inspector_dashboard', view='all'))
            images_query = images_query.filter_by(uploaded_by=user.id)

    # Paginar los resultados
    total_images = images_query.count()
    images = images_query.order_by(Image.upload_date.desc()).offset((page - 1) * IMAGES_PER_PAGE).limit(IMAGES_PER_PAGE).all()
    total_pages = (total_images + IMAGES_PER_PAGE - 1) // IMAGES_PER_PAGE

    return render_template(
        'inspector_dashboard.html',
        images=images,
        page=page,
        total_pages=total_pages,
        username=username,
        view=view
    )




@app.route('/')
def index():
    print(f"Rol del usuario en el template: {g.user_role}")
    return render_template('index.html')



def generar_informe_word_completo(filename, classification, confianza, inspector, informe_texto, graph_path):
    try:
        # Construir el nombre del archivo
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        word_filename = f"{filename.split('.')[0]}_informe_{timestamp}.docx"
        save_path = os.path.abspath(os.path.join(app.config['EXCEL_UPLOAD_FOLDER'], word_filename))

        # Verificar que el gráfico tenga una ruta completa
        absolute_graph_path = os.path.abspath(os.path.join('static', graph_path))

        # Crear el documento Word
        doc = Document()
        doc.add_heading("Informe Técnico de Análisis", level=1)
        doc.add_paragraph(f"Fecha: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Inspector: {inspector}")
        doc.add_paragraph(f"Nombre del archivo: {filename}")
        doc.add_paragraph(f"Clasificación: {classification}")
        doc.add_paragraph(" ")

        # Agregar gráfico si existe
        if os.path.exists(absolute_graph_path):
            doc.add_picture(absolute_graph_path, width=Inches(5.5))
        else:
            doc.add_paragraph("No se pudo incluir el gráfico.")
            print(f"El gráfico no se encontró en: {absolute_graph_path}")

        doc.add_heading('Informe Completo', level=2)
        doc.add_paragraph(informe_texto)

        # Guardar el archivo
        doc.save(save_path)
        print(f"Archivo Word guardado correctamente en: {save_path}")
        return word_filename
    except Exception as e:
        print(f"Error al generar el informe Word: {e}")
        return None



@app.route('/descargar_informe_word/<filename>', methods=['GET'])
@login_required
def descargar_informe_word(filename):
    try:
        ml_summary = session.get('ml_summary')
        anomaly_summary = session.get('anomaly_summary')
        graphs = session.get('graphs')
        informe_texto = session.get('informe_texto')
        punto_medicion = session.get('punto_medicion')  # Obtener el punto de medición
        eje_medido = session.get('eje_medido')          # Obtener el eje medido

        if not all([ml_summary, anomaly_summary, graphs, informe_texto, punto_medicion, eje_medido]):
            flash("Faltan datos para generar el informe.", "danger")
            return redirect(url_for('upload_excel'))

        # Crear documento Word
        doc = Document()
        doc.styles['Normal'].font.name = 'Arial'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        doc.styles['Normal'].font.size = Pt(11)

        # Encabezado
        doc.add_heading('Informe Técnico de Análisis', level=1)
        doc.add_paragraph(f"Fecha: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Inspector: {current_user.username}")
        doc.add_paragraph(f"Nombre del archivo: {filename}")
        doc.add_paragraph(f"Punto de medición: {punto_medicion}")
        doc.add_paragraph(f"Eje medido: {eje_medido}")

        # Clasificación General
        doc.add_heading('Clasificación General', level=2)
        table = doc.add_table(rows=1, cols=2, style='Table Grid')
        table.autofit = True
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Clasificación'
        hdr_cells[1].text = 'Cantidad'

        for key, value in ml_summary.items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(key)
            row_cells[1].text = str(value)

        # Clasificación de Anomalías
        doc.add_heading('Clasificación de Anomalías', level=2)
        for category, summary in anomaly_summary.items():
            doc.add_paragraph(category, style='Heading3')
            table = doc.add_table(rows=1, cols=2, style='Table Grid')
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Clasificación'
            hdr_cells[1].text = 'Cantidad'

            for cls, count in summary.items():
                row_cells = table.add_row().cells
                row_cells[0].text = str(cls)
                row_cells[1].text = str(count)

        # Gráficos Generados
        doc.add_heading('Gráficos Generados', level=2)
        for graph in graphs:
            doc.add_paragraph(graph['label'], style='Heading3')
            graph_path = os.path.join('static', graph['path'])
            if os.path.exists(graph_path):
                doc.add_picture(graph_path, width=Inches(5))

        # Informe Generado
        doc.add_heading('Informe Generado', level=2)
        for paragraph in informe_texto.split("\n\n"):
            doc.add_paragraph(paragraph)

        # Guardar en memoria y enviar
        word_file = io.BytesIO()
        doc.save(word_file)
        word_file.seek(0)

        return send_file(
            word_file,
            as_attachment=True,
            download_name=f"Informe_{filename.split('.')[0]}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        app.logger.error(f"Error al generar el informe Word: {str(e)}")
        flash("Error al generar el informe Word.", "danger")
        return redirect(url_for('upload_excel'))

app.config['GRAPH_UPLOAD_FOLDER'] = os.path.join(os.getcwd(), 'static', 'graphs')

if __name__ == '__main__':
    app.run(debug=True)
